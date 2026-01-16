"""
File Processing Utility for Chatbot
Extracts text and data from various file formats
"""

import os
import io
from typing import Optional, Dict, Any
from pathlib import Path
import mimetypes

# PDF processing
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word document processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Image processing with OCR
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# OpenAI for advanced image understanding
from openai import OpenAI
from config import get_settings

settings = get_settings()


class FileProcessor:
    """Process uploaded files and extract text content"""

    SUPPORTED_FORMATS = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'image/jpeg': 'image',
        'image/jpg': 'image',
        'image/png': 'image',
        'text/plain': 'text'
    }

    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    def __init__(self):
        self.openai_client = None
        if settings.openai_api_key:
            self.openai_client = OpenAI(api_key=settings.openai_api_key)

    def is_supported(self, content_type: str) -> bool:
        """Check if file type is supported"""
        return content_type in self.SUPPORTED_FORMATS

    def process_file(self, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """
        Process uploaded file and extract text

        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type of the file

        Returns:
            Dictionary with extracted text and metadata
        """
        if not self.is_supported(content_type):
            return {
                "success": False,
                "error": f"Unsupported file type: {content_type}. Supported types: PDF, DOCX, Images (JPG, PNG), TXT"
            }

        if len(file_content) > self.MAX_FILE_SIZE:
            return {
                "success": False,
                "error": f"File too large. Maximum size: {self.MAX_FILE_SIZE / 1024 / 1024}MB"
            }

        file_type = self.SUPPORTED_FORMATS[content_type]

        try:
            if file_type == 'pdf':
                extracted_text = self._extract_pdf(file_content)
            elif file_type == 'docx':
                extracted_text = self._extract_docx(file_content)
            elif file_type == 'image':
                extracted_text = self._extract_image(file_content, filename)
            elif file_type == 'text':
                extracted_text = self._extract_text(file_content)
            else:
                return {
                    "success": False,
                    "error": f"Handler not implemented for {file_type}"
                }

            if not extracted_text or extracted_text.strip() == "":
                return {
                    "success": False,
                    "error": "No text could be extracted from the file"
                }

            return {
                "success": True,
                "filename": filename,
                "file_type": file_type,
                "extracted_text": extracted_text,
                "text_length": len(extracted_text)
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Error processing file: {str(e)}"
            }

    def _extract_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            raise Exception("PDF processing not available. Install PyPDF2: pip install PyPDF2")

        pdf_file = io.BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)

        text = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)

        return "\n\n".join(text)

    def _extract_docx(self, file_content: bytes) -> str:
        """Extract text from Word document"""
        if not DOCX_AVAILABLE:
            raise Exception("DOCX processing not available. Install python-docx: pip install python-docx")

        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)

        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text.append(" | ".join(row_text))

        return "\n\n".join(text)

    def _extract_image(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from image using OCR or OpenAI Vision

        Priority: OpenAI Vision (more accurate) > Tesseract OCR
        """
        # Try OpenAI Vision first (more accurate for structured data)
        if self.openai_client:
            try:
                return self._extract_image_with_openai(file_content)
            except Exception as e:
                print(f"OpenAI Vision failed, falling back to OCR: {e}")

        # Fall back to Tesseract OCR
        if OCR_AVAILABLE:
            try:
                return self._extract_image_with_ocr(file_content)
            except Exception as e:
                raise Exception(f"OCR processing failed: {str(e)}")

        raise Exception("Image processing not available. Install Pillow and pytesseract, or configure OpenAI API key")

    def _extract_image_with_openai(self, file_content: bytes) -> str:
        """Extract text from image using OpenAI Vision API"""
        import base64

        # Convert image to base64
        base64_image = base64.b64encode(file_content).decode('utf-8')

        # Use OpenAI Vision to extract structured data
        response = self.openai_client.chat.completions.create(
            model="gpt-4o-mini",  # or "gpt-4-vision-preview"
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": """請從這張圖片中提取所有文字內容，特別注意：
                            - 公司名稱
                            - 產業別
                            - 資本額
                            - 專利數量（發明專利、新型專利）
                            - 認證資料
                            - 產品資訊（名稱、規格、價格等）

                            請以清晰的格式輸出所有找到的文字。"""
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=1000
        )

        return response.choices[0].message.content

    def _extract_image_with_ocr(self, file_content: bytes) -> str:
        """Extract text from image using Tesseract OCR"""
        image_file = io.BytesIO(file_content)
        image = Image.open(image_file)

        # Perform OCR
        # Add Chinese language support with eng+chi_tra (Traditional Chinese)
        try:
            text = pytesseract.image_to_string(image, lang='eng+chi_tra')
        except:
            # Fall back to English only if Chinese language pack not available
            text = pytesseract.image_to_string(image, lang='eng')

        return text

    def _extract_text(self, file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try UTF-8 first
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Try Big5 (Traditional Chinese encoding)
                return file_content.decode('big5')
            except UnicodeDecodeError:
                # Fall back to latin-1
                return file_content.decode('latin-1', errors='replace')


def get_file_info(filename: str, content_type: str) -> Dict[str, str]:
    """Get human-readable file information"""
    file_type_names = {
        'pdf': 'PDF文件',
        'docx': 'Word文件',
        'image': '圖片文件',
        'text': '文字文件'
    }

    processor = FileProcessor()
    file_type = processor.SUPPORTED_FORMATS.get(content_type, 'unknown')

    return {
        'filename': filename,
        'type': file_type_names.get(file_type, '未知格式'),
        'type_code': file_type
    }
